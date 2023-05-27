from docx import Document
import requests
from bs4 import BeautifulSoup
import openai
import feedparser
from dateutil.parser import parse
from datetime import datetime, timedelta
import pytz

openai.api_key = "YOUR_OPENAI_API_KEY"

# List of RSS feeds
rss_urls = [
    "https://news.xbox.com/en-us/feed/",
    "https://blog.playstation.com/feed/",
    "https://www.justice.gov/feeds/justice-news.xml?type%5Bpress_release%5D=press_release&type%5Bspeech%5D=speech&component%5B451%5D=451&&organization=1664&",
    "https://www.apple.com/tv-pr/news-feed.xml",
    "https://news.microsoft.com/category/press-releases/feed/",
    "https://www.reutersagency.com/feed/?taxonomy=best-topics&post_type=best",
]

entries = []

for rss_url in rss_urls:
    # Fetch the RSS feed
    feed = feedparser.parse(rss_url)

    # Filter entries from the last 24 hours
    time_limit = datetime.now(pytz.utc) - timedelta(hours=6)
    for entry in feed.entries:
        if "published" in entry:
            entry_date = parse(entry.published)
        elif "updated" in entry:
            entry_date = parse(entry.updated)
        else:
            continue

        if entry_date > time_limit:
            entries.append(entry)


for entry in entries:
    # Fetch the URL
    response = requests.get(entry.link)

    # Parse only the text within <h1> tags and <body> tags
    soup = BeautifulSoup(response.content, "html.parser")
    h1_text = soup.find("h1").get_text() if soup.find("h1") else ""
    body_text = soup.find("body").get_text() if soup.find("body") else ""

    page_content = h1_text + "\n" + body_text + "\n" + "Source: " + entry.link

    # The first part of the prompt you've created
    prompt_start = "I am going to send you text and you will rewrite it. You need to write it in a straightforward, human, and informative style. You will always rewrite the text as a news article for an audience whose age ranges from 17-25. You will always maintain AP style. You will always preserve the factual integrity of the original piece and include all main points.  You will start by creating a short eye-catching and exciting headline about the article. Include any quotes or statements from the original text. Focus on including any dates, especially release dates. Make the rewrite 600 words. You will use headers where appropriate. Each header will have at least two paragraphs and cover a substantial topic from the source material. After this, you will write an SEO Description AND a social media message to promote the article. The social media message will be a maximum of 255 characters. Here is the text to be rewritten: "

    # Adding the webpage content to your prompt
    full_prompt = prompt_start + page_content

    # Send the prompt to GPT-3
    response = openai.ChatCompletion.create(
        model="gpt-3.5-turbo",
        messages=[
            {"role": "system", "content": "You are an expert writer and editor."},
            {"role": "user", "content": full_prompt},
        ],
    )

    # Create a new Word Document
    doc = Document()

    # Add the GPT-3 response to the Word document
    doc.add_paragraph(response.choices[0].message["content"])

    # Add the URL at the end of the document
    doc.add_paragraph("\nSource: " + entry.link)  # this is the new line

    # Create a filename from the first 70 characters of the response
    # Remove or replace characters that are invalid in filenames
    invalid_chars = '/\?%*:|"<>. \n\r'
    filename = response.choices[0].message["content"][:70]
    for char in invalid_chars:
        filename = filename.replace(char, "_")

    # Save the document with the new filename
    doc.save(f"{filename}.docx")
